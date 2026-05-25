from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from openai import OpenAI
from dotenv import load_dotenv
from pydantic import BaseModel
import os
import shutil
import traceback
import re
from datetime import datetime

from docx import Document

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_community.vectorstores import FAISS

load_dotenv()

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://127.0.0.1:3000",
        "https://ai-chat-frontend-one.vercel.app",
        "https://ai-chat-frontend-h92larnka-jrolandomxs-projects.vercel.app",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

llm = ChatOpenAI(
    model="gpt-4.1-mini",
    api_key=os.getenv("OPENAI_API_KEY"),
)

vectorstore = None
uploaded_documents = []
last_article_review = ""

conversation_history = [
    {
        "role": "system",
        "content": (
            "Eres un asistente académico y técnico. "
            "Respondes de manera clara, breve y paso a paso."
        ),
    }
]


class ChatRequest(BaseModel):
    prompt: str


class TextRequest(BaseModel):
    text: str


class PDFQuestionRequest(BaseModel):
    question: str


@app.get("/")
def home():
    return {"message": "AI Academic Reviewer API funcionando correctamente"}


@app.post("/chat")
def chat(request: ChatRequest):
    conversation_history.append({"role": "user", "content": request.prompt})

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=conversation_history,
    )

    assistant_response = response.choices[0].message.content

    conversation_history.append(
        {"role": "assistant", "content": assistant_response}
    )

    return {
        "response": assistant_response,
        "history": conversation_history,
    }


@app.post("/chat-stream")
async def chat_stream(request: ChatRequest):
    conversation_history.append({"role": "user", "content": request.prompt})

    stream = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=conversation_history,
        stream=True,
    )

    async def generate():
        full_response = ""

        for chunk in stream:
            content = chunk.choices[0].delta.content

            if content:
                full_response += content
                yield content

        conversation_history.append(
            {"role": "assistant", "content": full_response}
        )

    return StreamingResponse(generate(), media_type="text/plain")


@app.post("/summarize")
def summarize(request: TextRequest):
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "Resume el texto en máximo 5 líneas."},
            {"role": "user", "content": request.text},
        ],
    )

    return {"summary": response.choices[0].message.content}


@app.post("/translate")
def translate(request: TextRequest):
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {"role": "system", "content": "Traduce el texto al inglés."},
            {"role": "user", "content": request.text},
        ],
    )

    return {"translation": response.choices[0].message.content}


@app.post("/keywords")
def keywords(request: TextRequest):
    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": "Extrae de 5 a 10 palabras clave del texto.",
            },
            {"role": "user", "content": request.text},
        ],
    )

    return {"keywords": response.choices[0].message.content}


@app.post("/upload-pdf")
def upload_pdf(file: UploadFile = File(...)):
    global vectorstore, uploaded_documents, last_article_review

    try:
        file_path = f"uploaded_{file.filename}"

        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)

        loader = PyPDFLoader(file_path)
        documents = loader.load()

        uploaded_documents = documents
        last_article_review = ""

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
        )

        chunks = text_splitter.split_documents(documents)

        embeddings = OpenAIEmbeddings(
            api_key=os.getenv("OPENAI_API_KEY"),
        )

        vectorstore = FAISS.from_documents(
            documents=chunks,
            embedding=embeddings,
        )

        return {
            "message": "PDF cargado correctamente",
            "filename": file.filename,
            "pages": len(documents),
            "chunks": len(chunks),
        }

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={"error": str(e)},
        )


@app.post("/ask-pdf")
def ask_pdf(request: PDFQuestionRequest):
    global vectorstore

    try:
        if vectorstore is None:
            return JSONResponse(
                status_code=400,
                content={"error": "Primero debes subir un PDF"},
            )

        results = vectorstore.similarity_search(request.question, k=3)

        context = "\n\n".join([doc.page_content for doc in results])

        prompt = f"""
Responde usando únicamente la información del contexto.

Si la respuesta no está en el contexto responde:
"No encontré esa información en el documento."

Contexto:
{context}

Pregunta:
{request.question}
"""

        response = llm.invoke(prompt)

        return {
            "question": request.question,
            "answer": response.content,
            "sources": [
                {
                    "page": doc.metadata.get("page"),
                    "content": doc.page_content[:300],
                }
                for doc in results
            ],
        }

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={"error": str(e)},
        )


def blind_review_text(text: str):
    patterns = [
        r"(?i)autor[a-z]*:.*",
        r"(?i)authors?:.*",
        r"(?i)afiliaci[oó]n:.*",
        r"(?i)universidad.*",
        r"(?i)correo.*",
        r"(?i)e-mail.*",
        r"(?i)email.*",
        r"(?i)orcid.*",
        r"(?i)agradecimientos.*",
    ]

    for pattern in patterns:
        text = re.sub(pattern, "[DATOS OCULTOS PARA REVISIÓN CIEGA]", text)

    return text


@app.post("/review-article")
def review_article(
    review_type: str = Form("Scopus"),
    blind_review: bool = Form(True),
):
    global uploaded_documents, last_article_review

    try:
        if not uploaded_documents:
            return JSONResponse(
                status_code=400,
                content={"error": "Primero debes subir un artículo PDF"},
            )

        full_text = "\n\n".join([doc.page_content for doc in uploaded_documents])

        if blind_review:
            full_text = blind_review_text(full_text)

        max_chars = 30000

        if len(full_text) > max_chars:
            full_text = full_text[:max_chars]

        review_prompt = f"""
Eres un sistema multiagente de arbitraje académico especializado en evaluación científica.

Tipo de evaluación:
{review_type}

Debes actuar simultáneamente como:

1. Revisor metodológico
2. Revisor teórico
3. Revisor editorial y de redacción
4. Revisor APA y formato
5. Editor en jefe

Tu función es elaborar un dictamen académico riguroso, crítico, objetivo,
constructivo y útil para mejorar la calidad científica, metodológica, teórica
y editorial del manuscrito.

Mantén un tono profesional, académico y humano. Evita frases genéricas,
superficiales o excesivamente duras.

Analiza el siguiente artículo científico:

ARTÍCULO:
{full_text}

Genera el dictamen con esta estructura exacta:

# Dictamen académico multiagente

## Badge de dictamen editorial

Indica SOLO UNA opción:
- Aceptado sin cambios
- Aceptado con cambios menores
- Requiere cambios mayores
- Rechazado

## Score general del artículo

Genera tabla markdown:

| Criterio | Score |
|---|---|
| Originalidad | X/10 |
| Metodología | X/10 |
| Marco teórico | X/10 |
| Redacción | X/10 |
| Resultados | X/10 |
| Discusión | X/10 |
| APA | X/10 |

## Revisión metodológica

Evalúa:
- claridad del problema de investigación;
- diseño metodológico;
- congruencia entre objetivos, metodología y resultados;
- tipo de estudio;
- población, muestra o corpus;
- técnicas de recolección de datos;
- técnicas de análisis;
- validez, confiabilidad o credibilidad;
- limitaciones;
- posibles sesgos;
- suficiencia de la explicación metodológica.

Incluye observaciones específicas y recomendaciones concretas.

## Revisión teórica

Evalúa:
- pertinencia del marco teórico;
- actualidad de la literatura;
- profundidad conceptual;
- claridad de categorías, variables o constructos;
- relación entre teoría, problema y hallazgos;
- originalidad;
- aporte científico;
- diálogo con literatura reciente.

Incluye observaciones específicas y recomendaciones concretas.

## Revisión editorial y de redacción

Evalúa:
- título;
- resumen;
- palabras clave;
- introducción;
- planteamiento del problema;
- objetivos;
- justificación;
- resultados;
- discusión;
- conclusiones;
- coherencia general;
- cohesión entre apartados;
- claridad argumentativa;
- redacción académica;
- repeticiones;
- ambigüedades;
- transiciones débiles;
- posible redacción artificial o excesivamente genérica.

Incluye observaciones específicas y recomendaciones concretas.

## Revisión APA y formato

Evalúa:
- uso de citas;
- correspondencia entre citas y referencias;
- suficiencia y actualidad de referencias;
- formato APA 7;
- tablas;
- figuras;
- mención de tablas y figuras en el texto;
- consistencia editorial.

No inventes referencias. Si no puedes verificar algo con el texto disponible, indícalo.

## Tabla sintética de observaciones

Genera tabla markdown:

| Apartado | Problema | Recomendación | Prioridad |
|---|---|---|---|

## Fortalezas

## Debilidades

## Recomendaciones concretas

## Evaluación por criterios

Genera tabla markdown:

| Criterio | Valoración cualitativa | Puntuación /10 |
|---|---|---|

Criterios:
- Originalidad
- Pertinencia del tema
- Rigor metodológico
- Sustento teórico
- Claridad de resultados
- Discusión
- Redacción académica
- Formato APA

## Dictamen final del editor en jefe

Incluye:
- Nivel de aporte científico
- Nivel de rigor metodológico
- Nivel de claridad editorial
- Recomendación editorial final, eligiendo solo una:
  - Aceptado sin cambios
  - Aceptado con cambios menores
  - Requiere cambios mayores
  - Rechazado

Redacta todo en español académico.
"""

        response = llm.invoke(review_prompt)

        last_article_review = response.content

        return {"review": last_article_review}

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={"error": str(e)},
        )


@app.post("/export-review")
def export_review():
    return export_review_word()


@app.post("/export-review-word")
def export_review_word():
    global last_article_review

    try:
        if not last_article_review:
            return JSONResponse(
                status_code=400,
                content={"error": "Primero debes generar un dictamen académico"},
            )

        document = Document()

        document.add_heading("Dictamen académico de artículo científico", level=1)

        document.add_paragraph(
            f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
        )

        document.add_paragraph("")

        for line in last_article_review.split("\n"):
            clean_line = line.strip()

            if not clean_line:
                document.add_paragraph("")
                continue

            if clean_line.startswith("# "):
                document.add_heading(clean_line.replace("# ", ""), level=1)

            elif clean_line.startswith("## "):
                document.add_heading(clean_line.replace("## ", ""), level=2)

            elif clean_line.startswith("### "):
                document.add_heading(clean_line.replace("### ", ""), level=3)

            elif clean_line.startswith("- "):
                document.add_paragraph(
                    clean_line.replace("- ", ""),
                    style="List Bullet",
                )

            else:
                document.add_paragraph(clean_line)

        file_path = "dictamen_academico.docx"

        document.save(file_path)

        return FileResponse(
            path=file_path,
            filename="dictamen_academico.docx",
            media_type=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
        )

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={"error": str(e)},
        )


@app.post("/export-review-pdf")
def export_review_pdf():
    global last_article_review

    try:
        if not last_article_review:
            return JSONResponse(
                status_code=400,
                content={"error": "Primero debes generar un dictamen académico"},
            )

        file_path = "dictamen_academico.pdf"

        doc = SimpleDocTemplate(file_path)
        styles = getSampleStyleSheet()
        elements = []

        elements.append(
            Paragraph("Dictamen académico de artículo científico", styles["Heading1"])
        )
        elements.append(Spacer(1, 12))

        elements.append(
            Paragraph(
                f"Fecha de generación: {datetime.now().strftime('%d/%m/%Y %H:%M')}",
                styles["BodyText"],
            )
        )
        elements.append(Spacer(1, 12))

        for line in last_article_review.split("\n"):
            clean_line = line.strip()

            if not clean_line:
                elements.append(Spacer(1, 8))
                continue

            clean_line = clean_line.replace("#", "")

            elements.append(
                Paragraph(clean_line, styles["BodyText"])
            )
            elements.append(Spacer(1, 8))

        doc.build(elements)

        return FileResponse(
            path=file_path,
            filename="dictamen_academico.pdf",
            media_type="application/pdf",
        )

    except Exception as e:
        traceback.print_exc()

        return JSONResponse(
            status_code=500,
            content={"error": str(e)},
        )